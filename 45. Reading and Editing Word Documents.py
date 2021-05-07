import docx
d = docx.Document('C:\\Users\\jenny\\OneDrive\\Documents\\Prof Development\\Python\\Automate the Boring Stuff\\demo.docx')
print(d.paragraphs)
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

p = d.paragraphs[1]
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

print(p.runs[1].bold)
print(p.runs[0].bold)
print(p.runs[0].bold == None)
print(p.runs[3].italic)

p.runs[3].underline == True
p.runs[3].text = 'italic and underline'
d.save('C:\\Users\\jenny\\OneDrive\\Documents\\Prof Development\\Python\\Automate the Boring Stuff\\demo2.docx')

print(p.style)
p.style ='Title'
d.save ('C:\\Users\\jenny\\OneDrive\\Documents\\Prof Development\\Python\\Automate the Boring Stuff\\demo3.docx')

d = docx.Document() # creating a blank new document
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('Hello this is another paragraph.')
d.save ('C:\\Users\\jenny\\OneDrive\\Documents\\Prof Development\\Python\\Automate the Boring Stuff\\demo4.docx')

p = d.paragraphs[0]
p.add_run('This is a new run.')
print(p.runs)
p.runs[1].bold = True
d.save ('C:\\Users\\jenny\\OneDrive\\Documents\\Prof Development\\Python\\Automate the Boring Stuff\\demo5.docx')

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('C:\\Users\\jenny\\OneDrive\\Documents\\Prof Development\\Python\\Automate the Boring Stuff\\demo.docx'))
